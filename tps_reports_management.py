import os, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "v5.settings")
django.setup()
from tps.models import *

import datetime
import locale
import subprocess
from io import BytesIO
locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches

def add_data(table, *args, color="EEEEEE"):
    size = len(table.rows)
    table.add_row()
    for index, arg in enumerate(args):
        table.cell(size, index).text = arg
        table.cell(size, index)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))

def add_images(document, images, width):
    img_table = document.add_table(rows=1, cols=len(images))
    for index, image in enumerate(images):
        paragraph = img_table.cell(0, index).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image, width=width)
    
def fix(s):
    s = str(s)
    if s == 'SCORE_Z': return 'Score Z'
    if s == 'TBL' or s == 'CBT': return s
    return ' '.join([o.capitalize() for o in s.split()]).replace('_', ' ').replace('De', 'de')

def get_month(m):
    return {
        1: 'Janeiro',
        2: 'Fevereiro',
        3: 'Março',
        4: 'Maio',
        5: 'Abril',
        6: 'Junho',
        7: 'Julho',
        8: 'Agosto',
        9: 'Setembro',
        10: 'Outubro',
        11: 'Novembro',
        12: 'Dezembro',
    }.get(m, '').upper()

START_DATE = datetime.datetime.strptime('01-08-2020', '%d-%m-%Y')
END_DATE = datetime.datetime.strptime('31-08-2020', '%d-%m-%Y')
CAMPUS = 'BSB'
GROUP = 'ENEM_ANUAL'
    
scores = TPSScore.objects.filter(campus=CAMPUS, group=GROUP, month=START_DATE.month).order_by('-score')
max_tps_answers = TPS.objects.filter(start_date__gte=START_DATE, end_date__lte=END_DATE, campus=CAMPUS, group=GROUP).count()
document = Document('RPA_management.docx')
table = document.tables[1]
scores_count = scores.count()
for score in scores:
    if not TPSAnswer.objects.filter(email=score.email).first():
        scores_count -= 1
        continue
    student_name = TPSAnswer.objects.filter(email=score.email).first().name
    student_answer_count = TPSAnswer.objects.filter(submission_date__gte=START_DATE, submission_date__lte=END_DATE, email=score.email, tps__campus=CAMPUS, tps__group=GROUP).count()
    add_data(table, '                         ' + student_name, str(score.score), '{} ({:.0f}%)'.format(student_answer_count, student_answer_count * 100.0 / max_tps_answers))

table2 = document.tables[0]
add_data(table2, '', 'Campus: ',  f'{CAMPUS}', color="FFFFFF")
add_data(table2, '', 'Grupo: ',  f'{GROUP}', color="FFFFFF")
add_data(table2, '', 'Mês: ', f'{get_month(START_DATE.month)}', color="FFFFFF")
add_data(table2, '', 'TPS Disponíveis: ', f'{max_tps_answers}', color="FFFFFF")
add_data(table2, '', 'Alunos: ', f'{scores_count}', color="FFFFFF")

document.save(f'reports/docx/27-08-2020/REL-ALUNOS-{CAMPUS}-{GROUP}-{get_month(START_DATE.month)}.docx')
subprocess.call(['libreoffice', '--headless', '--convert-to',  'pdf', f'reports/docx/27-08-2020/REL-ALUNOS-{CAMPUS}-{GROUP}-{get_month(START_DATE.month)}.docx', '--outdir', f'reports/pdf/27-08-2020/'])
    