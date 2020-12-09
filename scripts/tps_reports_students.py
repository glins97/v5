import os, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "v5.settings")
django.setup()
from tps.models import *

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import datetime
import locale
import subprocess
import plotly.express as px
from pandas import DataFrame
import plotly
from io import BytesIO
locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")

def add_data(table, *args, color="EEEEEE"):
    size = len(table.rows)
    table.add_row()
    for index, arg in enumerate(args):
        table.cell(size, index).text = arg
        table.cell(size, index)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))

def add_images(document, images, width):
    img_table = document.add_table(rows=1, cols=len(images))
    for index, image in enumerate(images):
        paragraph = img_table.cell(0, index).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image, width=width)
    
START_DATE = datetime.datetime.strptime('01/09/2020', '%d/%m/%Y')
END_DATE = datetime.datetime.strptime('30/09/2020', '%d/%m/%Y')
CAMPUS = 'BSB'
GROUP = 'PARTICULARES'
    
students = {
    tps_answer.email: 
    TPSAnswer.objects.filter(tps__start_date__gte=START_DATE, tps__end_date__lte=END_DATE, email=tps_answer.email, tps__campus=CAMPUS, tps__group=GROUP) for tps_answer in TPSAnswer.objects.filter(tps__start_date__gte=START_DATE, tps__end_date__lte=END_DATE, tps__campus=CAMPUS, tps__group=GROUP)
}

def get_subject(subject):
    subject = subject.lower()
    if 'bio' in subject:
        return 'Biologia'
    if 'mat' in subject:
        return 'Matemática'
    if 'fis' in subject or 'fís' in subject:
        return 'Física'
    if 'qui' in subject or 'quí' in subject:
        return 'Química'
    return 'Outros'

def fix(s):
    s = str(s)
    if s == 'SCORE_Z': return 'Score Z'
    if s == 'TBL' or s == 'CBT': return s
    return ' '.join([o.capitalize() for o in s.split()]).replace('_', ' ').replace('De', 'de')

scores = [obj.email for obj in TPSScore.objects.filter(month=START_DATE.month, campus=CAMPUS, group=GROUP).order_by('-score')]

for student in students:
    try:
        document = Document('RPA_students.docx')
        table = document.tables[1]
        grades = {
            'all': [],
            'bio': [],
            'fis': [],
            'qui': [],
            'mat': [],
        }
        # groups = {
        #     'SCORE_Z': 0,
        #     'TBL': 0,
        #     'CBT': 0,
        # }
        max_tps_answers = TPS.objects.filter(start_date__gte=START_DATE, end_date__lte=END_DATE, campus=CAMPUS, group=GROUP).count()
        if not max_tps_answers: continue
        rank = scores.index(student) + 1

        for answer in students[student]:
            if not answer.grade_group:
                continue
            add_data(table, '                         ' + get_subject(answer.tps.subject), str(answer.grade), fix(answer.grade_group), fix(answer.rank), fix(answer.submission_date.strftime('%d de %B')))
            grades['all'].append(answer.grade)
            # groups[answer.grade_group] += 1
            if 'bio' in answer.tps.subject.lower():
                grades['bio'].append(answer.grade)
            if 'qui' in answer.tps.subject.lower() or 'qui' in answer.tps.subject.lower():
                grades['qui'].append(answer.grade)
            if 'fis' in answer.tps.subject.lower() or 'fís' in answer.tps.subject.lower():
                grades['fis'].append(answer.grade)
            if 'mat' in answer.tps.subject.lower():
                grades['mat'].append(answer.grade)
            if 'qui' in answer.tps.subject.lower():
                grades['qui'].append(answer.grade)
        
        table2 = document.tables[0]
        add_data(table2, '', 'Aluno: ',  f'{answer.name}', color="FFFFFF")
        add_data(table2, '', 'Classificação geral: ', f'{rank}', color="FFFFFF")
        add_data(table2, '', 'Pontuação: ', '{}'.format(TPSScore.objects.filter(campus=CAMPUS, group=GROUP, month=START_DATE.month, email=answer.email).first().score), color="FFFFFF")
        add_data(table2, '', 'TPS Feitos: ', '{} ({:.0f}%)'.format(students[student].count(), students[student].count() * 100.0 / max_tps_answers), color="FFFFFF")
        add_data(table2, '', 'TPS Disponíveis: ',  f'{max_tps_answers}', color="FFFFFF")

        mat_png = BytesIO()
        bio_png = BytesIO()
        fis_png = BytesIO()
        qui_png = BytesIO()
        graphs = []
        
        if grades['all'] and len(grades['all']) > 1:
            all_png = BytesIO()
            px.line(DataFrame(grades['all'], columns=['Notas']), title="Notas Gerais", range_y=[0, 10], labels={
                "value": "Nota",
                "variable": "",
                "index": "TPS"
            }).write_image(all_png)
            graphs.append(all_png)
            
        if grades['fis'] and len(grades['fis']) > 1:
            fis_png = BytesIO()
            px.line(DataFrame(grades['fis'], columns=['Notas']), title="Notas de Física", range_y=[0, 10], labels={
                "value": "Nota",
                "variable": "",
                "index": "TPS"
            }).write_image(fis_png)
            graphs.append(fis_png)
        
        if grades['mat'] and len(grades['mat']) > 1:
            mat_png = BytesIO()
            px.line(DataFrame(grades['mat'], columns=['Notas']), title="Notas de Matemática", range_y=[0, 10], labels={
                "value": "Nota",
                "variable": "",
                "index": "TPS"
            }).write_image(mat_png)
            graphs.append(mat_png)

        if grades['bio'] and len(grades['bio']) > 1:
            bio_png = BytesIO()
            px.line(DataFrame(grades['bio'], columns=['Notas']), title="Notas de Biologia", range_y=[0, 10], labels={
                "value": "Nota",
                "variable": "",
                "index": "TPS"
            }).write_image(bio_png)
            graphs.append(bio_png)

        if grades['qui'] and len(grades['qui']) > 1:
            qui_png = BytesIO()
            px.line(DataFrame(grades['qui'], columns=['Notas']), title="Notas de Química", range_y=[0, 10], labels={
                "value": "Nota",
                "variable": "",
                "index": "TPS"
            }).write_image(qui_png)
            graphs.append(qui_png)

        if len(graphs) == 5:
            add_images(document, graphs[0:2], Inches(3.9))
            add_images(document, graphs[2:4], Inches(3.9))
            add_images(document, graphs[4:], Inches(3.9))

        if len(graphs) == 4:
            add_images(document, graphs[:2], Inches(3.9))
            add_images(document, graphs[2:], Inches(3.9))

        if len(graphs) == 3:
            add_images(document, graphs[0:2], Inches(3.9))
            add_images(document, graphs[2:], Inches(3.9))

        if len(graphs) == 2:
            add_images(document, graphs, Inches(3.9))

        if len(graphs) == 1:
            add_images(document, graphs, Inches(3.9))

        document.save(f'reports/docx/01-10-2020/{student}.docx')
        subprocess.call(['libreoffice', '--headless', '--convert-to',  'pdf', f'reports/docx/01-10-2020/{student}.docx', '--outdir', f'reports/pdf/01-10-2020/'])
        from mailer.mailer import send_templated_mail
        send_templated_mail('base.html', student, f'Relatório de Performance', f'reports/pdf/01-10-2020/{student}.pdf', title='Setembro de 2020', body=f'{answer.name.split()[0]}, confira seu desempenho neste mês, no relatório em anexo.<br>Lembre-se: o trabalho duro vence o talento.', footer="Equipe PPA")
    except Exception as e:
        print(repr(e))
        pass