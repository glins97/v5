import os, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "v5.settings")
django.setup()
from essay_manager.models import *

import datetime
import locale
import subprocess
from io import BytesIO
locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import numpy

def add_data(table, *args, color="EEEEEE"):
    size = len(table.rows)
    table.add_row()
    for index, arg in enumerate(args):
        table.cell(size, index).text = arg
        table.cell(size, index)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))

def add_images(document, images, width):
    img_table = document.add_table(rows=1, cols=len(images))
    for index, image in enumerate(images):
        paragraph = img_table.cell(0, index).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image, width=width)
    
def get_month(m):
    return {
        1: 'Janeiro',
        2: 'Fevereiro',
        3: 'Março',
        4: 'Maio',
        5: 'Abril',
        6: 'Junho',
        7: 'Julho',
        8: 'Agosto',
        9: 'Setembro',
        10: 'Outubro',
        11: 'Novembro',
        12: 'Dezembro',
    }.get(m, '').upper()

START_DATE = datetime.datetime.strptime('01-09-2020', '%d-%m-%Y')
END_DATE = datetime.datetime.strptime('30-09-2020', '%d-%m-%Y')
JURY = 'VUNESP'
    
essays = Essay.objects.filter(grade__gt=0, theme__jury=JURY, upload_date__gte=START_DATE, upload_date__lte=END_DATE).exclude(theme__description="Solidário")
students = {essay.user: Essay.objects.filter(grade__gt=0, user=essay.user, theme__jury=JURY, upload_date__gte=START_DATE, upload_date__lte=END_DATE) for essay in essays}

document = Document('RPA_ESSAYS_management.docx')
table = document.tables[1]
for student in students:
    student_essays_count = len(students[student])
    student_essays_average_grade = int(numpy.mean([essay.grade for essay in students[student]]))
    add_data(table, '                         ' + f'{student.first_name} {student.last_name}', f'{student_essays_count}', f'{student_essays_average_grade}')

average_grade = int(numpy.mean([essay.grade for essay in essays]))

table2 = document.tables[0]
add_data(table2, '', 'Banca: ',  f'{JURY}', color="FFFFFF")
add_data(table2, '', 'Alunos: ',  f'{len(students)}', color="FFFFFF")
add_data(table2, '', 'Redações: ',  f'{len(essays)}', color="FFFFFF")
add_data(table2, '', 'Nota média: ',  f'{average_grade}', color="FFFFFF")

document.save(f'reports/docx/01-10-2020/REL-ESSAYS-{JURY}-{get_month(START_DATE.month)}.docx')
subprocess.call(['libreoffice', '--headless', '--convert-to',  'pdf', f'reports/docx/01-10-2020/REL-ESSAYS-{JURY}-{get_month(START_DATE.month)}.docx', '--outdir', f'reports/pdf/01-10-2020/'])
    