import os, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "v5.settings")
django.setup()
from essay_manager.models import *

import datetime
import locale
import subprocess
from io import BytesIO
locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import numpy
import json
import plotly.express as px
from pandas import DataFrame

def add_data(table, *args, color="EEEEEE"):
    size = len(table.rows)
    table.add_row()
    for index, arg in enumerate(args):
        table.cell(size, index).text = arg
        table.cell(size, index)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))

def add_images(document, images, width):
    img_table = document.add_table(rows=1, cols=len(images))
    for index, image in enumerate(images):
        paragraph = img_table.cell(0, index).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image, width=width)
    
def get_month(m):
    return {
        1: 'Janeiro',
        2: 'Fevereiro',
        3: 'Março',
        4: 'Maio',
        5: 'Abril',
        6: 'Junho',
        7: 'Julho',
        8: 'Agosto',
        9: 'Setembro',
        10: 'Outubro',
        11: 'Novembro',
        12: 'Dezembro',
    }.get(m, '')

START_DATE = datetime.datetime.strptime('01-09-2020', '%d-%m-%Y')
END_DATE = datetime.datetime.strptime('30-09-2020', '%d-%m-%Y')
JURY = 'VUNESP'
    
essays = Essay.objects.filter(grade__gt=0, theme__jury=JURY, upload_date__gte=START_DATE, upload_date__lte=END_DATE).exclude(theme__description="Solidário")
students = {essay.user: Essay.objects.filter(grade__gt=0, user=essay.user, theme__jury=JURY, upload_date__gte=START_DATE, upload_date__lte=END_DATE).exclude(theme__description="Solidário") for essay in essays}
for student in students:
    document = Document('RPA_ESSAYS_students.docx')
    table = document.tables[1]
    student_essays_count = len(students[student])

    student_essays_average_grade = int(numpy.mean([essay.grade for essay in students[student]]))
    for essay in students[student]:
        add_data(table, '', f'{essay.theme.description[0].upper() + essay.theme.description[1:]}', f'{essay.grade}', f'{essay.upload_date.strftime("%d de %B")}')
    
    print(student.username, student_essays_count)
    if student_essays_count > 1:
        competencies = [json.loads(Correction.objects.get(essay=essay).data)['competencies']['grades'] for essay in students[student]]
        if JURY == 'ENEM':
            graphs = []
            for comp in range(1, 6):
                out = BytesIO()
                fig = px.line(DataFrame([comps[f'a{comp}'] for comps in competencies], columns=[f'Competência {comp}']), title=f'Competência {comp}', range_y=[40, 210], labels={
                    "value": "Nota",
                    "variable": "",
                    "index": "Redação"
                })
                fig.update_yaxes(tick0=0, dtick=40)
                fig.write_image(out)
                graphs.append(out)

            add_images(document, graphs[0:2], Inches(3.9))
            add_images(document, graphs[2:4], Inches(3.9))
            add_images(document, graphs[4:], Inches(3.9))
        elif JURY == 'VUNESP':
            graphs = []
            out_a = BytesIO()
            fig = px.line(DataFrame([comps[f'a'] for comps in competencies], columns=[f'Critério A']), title=f'Critério A', range_y=[0, 3], labels={
                "value": "Nota",
                "variable": "",
                "index": "Redação"
            })
            fig.update_yaxes(tick0=0, dtick=1)
            fig.write_image(out_a)
            graphs.append(out_a)
            
            out_b = BytesIO()
            fig = px.line(DataFrame([comps[f'b'] for comps in competencies], columns=[f'Critério B']), title=f'Critério B', range_y=[0, 4], labels={
                "value": "Nota",
                "variable": "",
                "index": "Redação"
            })
            fig.update_yaxes(tick0=0, dtick=1)
            fig.write_image(out_b)
            graphs.append(out_b)
            
            out_c = BytesIO()
            fig = px.line(DataFrame([comps[f'c'] for comps in competencies], columns=[f'Critério C']), title=f'Critério C', range_y=[0, 4], labels={
                "value": "Nota",
                "variable": "",
                "index": "Redação"
            })
            fig.update_yaxes(tick0=0, dtick=1)
            fig.write_image(out_c)
            graphs.append(out_c)
            
    table2 = document.tables[0]
    add_data(table2, '', 'Aluno: ',  f'{student.first_name} {student.last_name}', color="FFFFFF")
    add_data(table2, '', 'Mês: ', f'{get_month(START_DATE.month)} de {START_DATE.year}', color="FFFFFF")
    add_data(table2, '', 'Banca: ', f'{JURY}', color="FFFFFF")
    add_data(table2, '', 'Redações: ', f'{student_essays_count}', color="FFFFFF")
    add_data(table2, '', 'Nota média: ', f'{student_essays_average_grade}', color="FFFFFF")

    document.save(f'reports/docx/01-10-2020/{student.username}.docx')
    subprocess.call(['libreoffice', '--headless', '--convert-to',  'pdf', f'reports/docx/01-10-2020/{student.username}.docx', '--outdir', f'reports/pdf/01-10-2020/'])
    # from mailer.mailer import send_templated_mail
    # send_templated_mail('base.html', student.username, f'Relatório de Performance em Redações', f'reports/pdf/01-10-2020/{student.username}.pdf', title='Setembro de 2020', body=f'{student.first_name.capitalize()}, confira seu desempenho em setembro na produção de redações, no relatório em anexo.<br><br>"Somos aquilo que fazemos repetidamente. A excelência, portanto, não é um ato, mas um hábito" - Aristóteles', footer="Equipe PPA")
